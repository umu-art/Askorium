"""
Unit tests for document_renderer module.

Tests the pure rendering functions (_render_docx, _render_xlsx, _render_pdf)
by constructing in-memory document bytes with the same libraries used in production.
"""
import io
import pytest

from app.document_renderer import _render_docx, _render_xlsx, DocumentRenderer


class TestRenderDocx:
    def _make_docx(self, paragraphs: list[tuple[str, str]]) -> bytes:
        """Create a DOCX in memory. Each item is (text, style_name)."""
        from docx import Document
        doc = Document()
        for text, style in paragraphs:
            try:
                para = doc.add_paragraph(text, style=style)
            except Exception:
                para = doc.add_paragraph(text)
                if style:
                    para.style = doc.styles["Normal"]
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_normal_paragraph(self):
        data = self._make_docx([("Hello world", "Normal")])
        html = _render_docx(data)
        assert "<p>Hello world</p>" in html
        assert html.startswith("<html>")

    def test_heading_paragraph(self):
        data = self._make_docx([("Chapter 1", "Heading 1")])
        html = _render_docx(data)
        assert "<h1>Chapter 1</h1>" in html

    def test_heading_level_2(self):
        data = self._make_docx([("Section", "Heading 2")])
        html = _render_docx(data)
        assert "<h2>Section</h2>" in html

    def test_multiple_paragraphs_order(self):
        data = self._make_docx([
            ("First", "Normal"),
            ("Second", "Normal"),
        ])
        html = _render_docx(data)
        assert html.index("First") < html.index("Second")

    def test_empty_paragraphs_skipped(self):
        data = self._make_docx([
            ("", "Normal"),
            ("Non-empty", "Normal"),
        ])
        html = _render_docx(data)
        assert html.count("<p>") == 1
        assert "Non-empty" in html

    def test_mixed_content(self):
        data = self._make_docx([
            ("Title", "Heading 1"),
            ("Body text", "Normal"),
        ])
        html = _render_docx(data)
        assert "<h1>Title</h1>" in html
        assert "<p>Body text</p>" in html

    def test_output_wrapped_in_html_body(self):
        data = self._make_docx([("text", "Normal")])
        html = _render_docx(data)
        assert html.startswith("<html><body>")
        assert html.endswith("</body></html>")


class TestRenderXlsx:
    def _make_xlsx(self, sheets: dict[str, list[list]]) -> bytes:
        """Create an XLSX in memory. sheets maps sheet_name → rows of values."""
        from openpyxl import Workbook
        wb = Workbook()
        first = True
        for name, rows in sheets.items():
            if first:
                ws = wb.active
                ws.title = name
                first = False
            else:
                ws = wb.create_sheet(name)
            for row in rows:
                ws.append(row)
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def test_single_sheet_title(self):
        data = self._make_xlsx({"Sheet1": [["A", "B"]]})
        html = _render_xlsx(data)
        assert "<h2>Sheet1</h2>" in html

    def test_cell_values_rendered(self):
        data = self._make_xlsx({"Data": [["Name", "Age"], ["Alice", 30]]})
        html = _render_xlsx(data)
        assert "<td>Name</td>" in html
        assert "<td>Alice</td>" in html
        assert "<td>30</td>" in html

    def test_none_cells_rendered_empty(self):
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "S"
        ws.append(["val", None])
        buf = io.BytesIO()
        wb.save(buf)
        data = buf.getvalue()

        html = _render_xlsx(data)
        assert "<td></td>" in html

    def test_multiple_sheets(self):
        data = self._make_xlsx({
            "Alpha": [["a"]],
            "Beta": [["b"]],
        })
        html = _render_xlsx(data)
        assert "<h2>Alpha</h2>" in html
        assert "<h2>Beta</h2>" in html

    def test_output_wrapped_in_html_body(self):
        data = self._make_xlsx({"S": [["x"]]})
        html = _render_xlsx(data)
        assert html.startswith("<html><body>")
        assert html.endswith("</body></html>")

    def test_rows_wrapped_in_tr(self):
        data = self._make_xlsx({"S": [["v1", "v2"]]})
        html = _render_xlsx(data)
        assert "<tr>" in html
        assert "</tr>" in html


class TestDocumentRendererDispatch:
    @pytest.mark.asyncio
    async def test_unsupported_hint_raises_value_error(self):
        renderer = DocumentRenderer()
        with pytest.raises(ValueError, match="unsupported document type"):
            await _call_render_with_fake_download(renderer, b"data", "pptx")

    @pytest.mark.asyncio
    async def test_dispatch_pdf(self, monkeypatch):
        """render() with hint='pdf' calls _render_pdf, not docx/xlsx."""
        import app.document_renderer as mod

        called_with = {}

        def fake_pdf(data):
            called_with["type"] = "pdf"
            called_with["data"] = data
            return "<html><body><p>pdf</p></body></html>"

        monkeypatch.setattr(mod, "_render_pdf", fake_pdf)

        renderer = DocumentRenderer()
        monkeypatch.setattr(renderer, "_download", _make_fake_download(b"pdfbytes"))

        result = await renderer.render("https://example.com/doc.pdf", "pdf")
        assert called_with["type"] == "pdf"
        assert called_with["data"] == b"pdfbytes"
        assert result == "<html><body><p>pdf</p></body></html>"

    @pytest.mark.asyncio
    async def test_dispatch_docx(self, monkeypatch):
        import app.document_renderer as mod

        called_with = {}

        def fake_docx(data):
            called_with["type"] = "docx"
            return "<html><body></body></html>"

        monkeypatch.setattr(mod, "_render_docx", fake_docx)

        renderer = DocumentRenderer()
        monkeypatch.setattr(renderer, "_download", _make_fake_download(b"docxbytes"))

        await renderer.render("https://example.com/file.docx", "docx")
        assert called_with["type"] == "docx"

    @pytest.mark.asyncio
    async def test_dispatch_xlsx(self, monkeypatch):
        import app.document_renderer as mod

        called_with = {}

        def fake_xlsx(data):
            called_with["type"] = "xlsx"
            return "<html><body></body></html>"

        monkeypatch.setattr(mod, "_render_xlsx", fake_xlsx)

        renderer = DocumentRenderer()
        monkeypatch.setattr(renderer, "_download", _make_fake_download(b"xlsxbytes"))

        await renderer.render("https://example.com/file.xlsx", "xlsx")
        assert called_with["type"] == "xlsx"

    @pytest.mark.asyncio
    async def test_hint_case_insensitive(self, monkeypatch):
        """Hint should be lowercased before dispatch."""
        import app.document_renderer as mod

        called = {}

        monkeypatch.setattr(mod, "_render_xlsx", lambda d: (called.update({"ok": True}), "<html/>")[1])
        renderer = DocumentRenderer()
        monkeypatch.setattr(renderer, "_download", _make_fake_download(b"x"))

        await renderer.render("https://example.com/file.XLSX", "XLSX")
        assert called.get("ok") is True


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_fake_download(content: bytes):
    async def _download(url: str) -> bytes:
        return content
    return _download


async def _call_render_with_fake_download(renderer: DocumentRenderer, data: bytes, hint: str):
    renderer._download = _make_fake_download(data)
    return await renderer.render("https://example.com/file", hint)
